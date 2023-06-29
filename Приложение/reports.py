import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
import numpy as np
from io import BytesIO 
import os
import matplotlib.colors as mcolors
import random
class Reports ():  
#  запись строки в отчет
    def print_row (document, row):
            add_p = document.add_paragraph()
            r = add_p.add_run(row)
            r.font.name = 'TimesNewRoman'
            r.font.size = Pt(9)
            p_format = add_p.paragraph_format
            p_format.space_before = Pt(1)
            p_format.space_after = Pt(1)
# добавить строку в таблицу
    def add_row_in_table(table, row):
        for i, item in enumerate(row):
            table.rows[len(table.rows)-1].cells[i].text = item
    def random_color_generator(n):
        colors = []
        for i in range(n):
            colors.append(random.choice(list(mcolors.CSS4_COLORS.keys())))
        return colors
#структура показателя на уровне фо->субъект рф
    def pie_F (data, lists, years, parameters):
        document = Document()        
        regions, fds = lists
        values = []
        trash = []  
        
        for param in parameters:
            for fd in fds:
                for year in years:
                    table = document.add_table(1, 4)
                    header = ['Федеральный округ', 'Субъект РФ', 'Год', f'{param}']
                    Reports.add_row_in_table(table, header)
                    if len(regions)==0:
                        regions = data['Субъект_РФ'].loc[(data['ФО'] == fd) & (data['Год'] == year)].sort_values()
                    else:
                        regions = data['Субъект_РФ'].loc[(data['Субъект_РФ'].isin(regions))&(data['Год'] == year)].sort_values()
                    for region in regions:
                        value = data[param].loc[(data['Субъект_РФ']==region) & (data['Год']==year)].fillna(0).values[0]
                        if int(value) == 0:
                            trash.append(regions[regions==region].index.values[0])
                        row = [fd, region, str(year), str(value)]
                        
                        table.add_row()
                        Reports.add_row_in_table(table, row)
                        values.append(value)
                    total = sum(values)
                    row = ['Сумма', str(total)]
                    table.add_row()
                    for i, item in enumerate(row):
                        table.rows[len(table.rows)-1].cells[2+i].text = item
                    for i in range(len(trash)): values.remove(0)
                    if len(trash) != 0 :
                        regions = regions.drop(labels= trash)
                        trash = []
                    values_norm = [i/total for i in values]
                    memfile = BytesIO()
                    plt.title(f'{fd}, {year} : {param}', pad = 200)
                    plt.pie(values_norm, labels = regions, autopct='%1.1f%%', rotatelabels=True)
                    plt.gcf().set_size_inches(5, 5)
                    plt.savefig(memfile, bbox_inches = 'tight')
                    plt.close('all')
                    document.add_picture(memfile, width=Inches(5))
                    memfile.close()
                    d = dict(zip(regions, values_norm))
                    table = document.add_table(0, 3)
                    for region, value in d.items():
                        table.add_row()
                        row = [region, str(year), f'{f"{value*100:1.1f}%":>5s}']
                        Reports.add_row_in_table(table, row)
                    document.add_paragraph('\n')
                    values = []
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        document.save(desktop+'\Структура показателя на уровне ФО.docx')
# список фо->субъекты рф = статистические показатели с результирующей суммой для федерального округа за год. списки организованы в рамках фо -> год  
    def list_FO (data, lists,  years, parameters):
        document = Document()
        table = document.add_table(1, 3)
        regions, fds = lists
        summers = []
        total = []
        for fd in fds:
            header = ['Федеральный округ', 'Год', 'Субъект РФ']
            for param in parameters:
                total.append([])
                header.append(param)
                table.add_column(Inches(1))
            Reports.add_row_in_table(table, header)
            table.add_row()
            table.rows[len(table.rows)-1].cells[0].text = f'{fd}:'
            for Год in years:
                for param in parameters:
                    summers.append([])
                table.add_row()
                table.rows[len(table.rows)-1].cells[1].text = f'{Год:>d}'
                if len(regions) == 0 :
                    regions = data['Субъект_РФ'].loc[(data['ФО'] == fd) & (data['Год'] == Год)].sort_values()
                else:
                    regions = data['Субъект_РФ'].loc[(data['Субъект_РФ'].isin(regions))& (data['Год'] == Год)].sort_values()
                for region in regions:
                    table.add_row()
                    table.rows[len(table.rows)-1].cells[2].text = region
                    for i in range(len(parameters)):
                        value = data[parameters[i]].loc[(data['Субъект_РФ']==region) & (data['Год']==Год)].values[0]
                        table.rows[len(table.rows)-1].cells[3+i].text = f"{f'{value:.2f}':>10s}"
                        summers[i].append(value)
                table.add_row()
                table.rows[len(table.rows)-1].cells[2].text = 'Сумма'
                for i in range(len(parameters)):
                    summers[i] = np.nan_to_num(summers[i])
                    table.rows[len(table.rows)-1].cells[3+i].text =f"{f'{sum(summers[i]):.2f}':>10s}" 
                    total[i].append(sum(summers[i]))
                summers = []
            table.add_row()
            table.rows[len(table.rows)-1].cells[1].text = 'Сумма'
            for i, param in enumerate(parameters): 
                    summers.append([])
                    table.rows[len(table.rows)-1].cells[3+i].text = f'{f"{sum(total[i]):.2f}":>10s}'
            total = []
            document.add_paragraph('\n')
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        document.save(desktop+'\Субъекты РФ и стат.показатели.docx')
 # рейтинг субъектов рф в рамках ФО за 1 год по стат.показателям с гистограммами
    def rating_per_param (data, lists, years, parameters):
        document = Document()
        regions, fds = lists
        labels = []
        for fd in fds:
            Reports.print_row(document,f'{fd:^100}' )
            for year in years:
                for param in parameters:
                    table = document.add_table(1, 2)
                    header = ['№', 'Субъект РФ']
                    table.add_column(width = Inches(1))
                    header.append(f" {param:>{len(param)}}")
                    header.append(str(year))
                    Reports.add_row_in_table(table, header)
                    if len(regions)==0:
                        values = data[param].loc[(data['ФО'] == fd) & (data['Год'] == year)].sort_values(ascending = False)
                    else:
                        values = data[param].loc[(data['Субъект_РФ'].isin(regions))&(data['Год'] == year)].sort_values(ascending = False)
                    values = values[np.isfinite(values)]
                    for i, value in enumerate(values):
                        if np.isnan([value]): break
                        subject = data['Субъект_РФ'].loc[(data[param] == value) & (data['Год'] == year)].values[0]
                        labels.append(f"{i} {subject}")
                        row = [f"{i+1:>2d}", str(subject), f'{value:.1f}']
                        table.add_row()
                        Reports.add_row_in_table(table, row)
                    values = [int(item) for item in values]
                    memfile = BytesIO()
                    fig, ax = plt.subplots()
                    ax.bar(np.arange(len(values)), values, color = Reports.random_color_generator(len(values)), label = labels)
                    ax.set_xticks([i for i in range(0, len(values))])
                    box = ax.get_position()
                    ax.set_position([box.x0, box.y0, box.width * 0.8, box.height])
                    ax.legend(loc='upper left', bbox_to_anchor=(1, 1))
                    plt.gcf().set_size_inches(5, 5)
                    plt.savefig(memfile, bbox_inches = 'tight')
                    document.add_picture(memfile, width=Inches(5))
                    plt.close('all')
                    memfile.close()
                    labels = []
                    document.add_paragraph('\n\n')
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        document.save(desktop+'\Рейтинг субъектов РФ в рамках ФО и года.docx')
# рейтинг субъектов РФ в рамках 1 года по показателям
    def rating (data, regions, years, parameters):
        document = Document()
        for param in parameters:
            for Год in years:
                table = document.add_table(1, 5)
                values = data[param].loc[(data['Год']==Год)].sort_values(ascending=False)
                values_index = values.index
                header  =['№', 'Субъект РФ', f'{param:^{len(param)}s}', 'Федеральный округ', f'Год:   {Год}']
                Reports.add_row_in_table(table, header)
                i=0
                for index in values_index:
                    subject = data['Субъект_РФ'].iloc[index]
                    if subject in regions:
                        if i != len(values): i +=1
                        fd = data['ФО'].loc[(data['Субъект_РФ'] == subject)].values[0]
                        row = [f'{i:>2d}', str(subject), f'{data[param].iloc[index]}', fd]
                        table.add_row()
                        Reports.add_row_in_table(table, row)
                    else: 
                        continue
                document.add_paragraph('\n\n')
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        document.save(desktop+'\Рейтинг субъектов РФ в рамках года.docx')
 #  список значений стат.показателей для субъектов рф (алфавитный порядок) за все года с рядами динамики данных показателей
    def list_and_dynamic_plot (data, regions, years, parameters):
        document = Document()
        header = ['Субъект РФ', 'Федеральный округ', 'Год']
        for param in parameters:
            header.append(f"{param}")
        for region in regions:
            table = document.add_table(1, 3)
            for param in parameters:
                table.add_column(Inches(1))
            Reports.add_row_in_table(table, header)
            for Год in years:
                    val =  data.loc[(data['Субъект_РФ']==region)&(data['Год'] == Год)].values[0]
                    row = [str(val[0]), str(val[1]), str(val[2])]
                    for i, param in enumerate(parameters):
                        row.append(f'{val[3+i]:.1f}')
                    table.add_row()
                    Reports.add_row_in_table(table, row)
            values = []
            memfile = BytesIO()
            fig = plt.figure()
            fig.add_subplot(111, frame_on = False)
            plt.tick_params(labelcolor='none', bottom = False, left = False)
            plt.tight_layout()
            for i, param in enumerate(parameters):
                values = data[param].loc[(data['Субъект_РФ']==region) & (data['Год'].isin(years))].values
                ax = fig.add_subplot(1, len(parameters), i+1)
                ax.plot(years, values, marker = '.')
                ax.title.set_text(param)
                values = []
            plt.subplots_adjust(left=None, bottom=None, right=None, top=None, wspace=1, hspace=None)
            plt.gcf().set_size_inches(5, 5)
            plt.savefig(memfile, bbox_inches = 'tight')
            document.add_picture(memfile, width=Inches(5) )
            plt.close('all')
            memfile.close
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        path = desktop+'\Стат.показатели с динамикой'
        document.save(path+'.docx')
        
        
        